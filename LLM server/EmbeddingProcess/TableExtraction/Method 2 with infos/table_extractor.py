# -*- coding: utf8 -*-

import csv
import xlwt
import openpyxl
import json
import datetime

from docx import Document
from docx.table import _Cell
from docx.oxml.simpletypes import ST_Merge


def __extract_table(table, strip_space=False):
    """Extracts table data from table object"""
    results = []
    n = 0
    for tr in table._tbl.tr_lst:
        r = []
        for tc in tr.tc_lst:
            for grid_span_idx in range(tc.grid_span):
                if tc.vMerge == ST_Merge.CONTINUE:
                    value = results[n - 1][len(r) - 1]
                elif grid_span_idx > 0:
                    value = r[-1]
                else:
                    cell = _Cell(tc, table)
                    # Normalize and remove unwanted spaces
                    cell.text = repr(cell.text)
                    value = cell.text.replace("\\xa0", " ")  # Convert non-breaking spaces to normal spaces
                    value = value.replace("\n", " ")  # Remove unwanted newlines

                if strip_space:
                    value = value.strip()
                r.append(value)
        results.append(r)
        n += 1
    return results


def __store_table(tabdata, filename, format="csv"):
    """Saves table data as csv file"""
    if format == "csv":
        f = open(filename, "w", encoding='utf8')
        w = csv.writer(f, delimiter=",")
        for row in tabdata:
            w.writerow(row)
    elif format == 'tsv':
        f = open(filename, 'w')
        w = csv.writer(f, delimiter='\t')
        for row in tabdata:
            w.writerow(row)
    elif format == 'xls':
        workbook = xlwt.Workbook()
        ws = __xls_table_to_sheet(tabdata, workbook.add_sheet("0"))
#        print(dir(ws))
        workbook.save(filename)
    elif format == "xlsx":
        workbook = openpyxl.Workbook()
        ws = __xlsx_table_to_sheet(tabdata, workbook.create_sheet("0"))
        workbook.save(filename)

def __xls_table_to_sheet(table, ws):
    rn = 0
    for row in table:
        cn = 0
        for c in row:
            ws.write(rn, cn, c)
            cn += 1
        rn += 1
    return ws


def __xlsx_table_to_sheet(table, ws):
    rn = 0
    for row in table:
        ws.append(row)
        rn += 1
    return ws


def extract_tables(filename, strip_space=True):
    """Extracts table from .DOCX files"""
    tables = []
    document = Document(filename)
    n = 0
    for table in document.tables:
        n += 1
        info = {}
        info['id'] = n
        info['num_cols'] = len(table.columns)
        info['num_rows'] = len(table.rows)
        info['style'] = table.style.name
        tdata = __extract_table(table, strip_space=strip_space)
        info['data'] = tdata
        tables.append(info)
    return tables





def extract(filename, format="csv", sizefilter=0, singlefile=False, output=None, strip_space=True):
    """Extracts tables from csv files and saves them as csv, xls or xlsx files"""
    tables = extract_tables(filename, strip_space=strip_space)
    name = filename.rsplit(".", 1)[0]
    format = format.lower()
    n = 0
    lfilter = int(sizefilter)
    if singlefile:
        if format == "xls":
            workbook = xlwt.Workbook()
            for t in tables:
                if lfilter >= len(t):
                    continue
                n += 1
                ws = __xls_table_to_sheet(t['data'], workbook.add_sheet(str(n)))
            destname = output if output else name + ".%s" % (format)
            workbook.save(destname)
        elif format == "xlsx":
            workbook = openpyxl.Workbook()
            for t in tables:
                if lfilter >= len(t):
                    continue
                n += 1
                ws = __xlsx_table_to_sheet(t['data'], workbook.create_sheet(str(n)))
            destname = output if output else name + ".%s" % (format)
            workbook.save(destname)
        elif format == "json":
            report = {'filename' : filename,
            'timestamp' : datetime.datetime.now().isoformat(), 'num_tables' : len(tables),
            'tables' : tables}
            destname = output if output else name + ".%s" % (format)
            f = open(destname, 'w', encoding='utf8')
            json.dump(report, f, ensure_ascii=False, indent=4)
            f.close()

    else:
        for t in tables:
            if lfilter >= len(t):
                continue
            n += 1
            destname = output if output else name + "_%d.%s" % (n, format)
            __store_table(t['data'], destname, format)


def analyze(filename):
    """Analyzes docx file"""
    tableinfo = []
    document = Document(filename)
    n = 0
    for table in document.tables:
        n += 1
        info = {}
        info['id'] = n
        info['num_cols'] = len(table.columns)
        info['num_rows'] = len(table.rows)
        info['style'] = table.style.name
        tableinfo.append(info)
    return tableinfo




docx_file = "../Testing Docx/22104-i30.docx"
output_file = ".\\tables_testing"

print(analyze(docx_file))
tables = extract_tables(filename=docx_file)

with open(f'{output_file}\\22104-i30.json', 'w', encoding='utf-8') as f:
    json.dump(dict(enumerate(tables)), f, ensure_ascii=False, indent=4)


